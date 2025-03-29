import streamlit as st
import io
import random
import re
import string
import web_scraper
import os
from urllib.parse import urljoin, urlparse
from docx import Document
import pandas as pd

# Set an approximate maximum token limit (using words as a proxy)
MAX_TOKEN_LIMIT = 15000

def normalize_url(url):
    """Ensure the URL begins with 'https://'. Replace 'http://' with 'https://' or add the prefix if missing."""
    url = url.strip()
    if not url:
        return None
    if url.startswith("http://"):
        url = "https://" + url[len("http://") :]
    elif not url.startswith("https://"):
        url = "https://" + url
    return url

def generate_random_filename(url):
    """
    Generate a base filename of exactly 15 alphanumeric characters.
    Take the first 8 alphanumeric characters from the URL (after 'https://')
    and append random alphanumeric characters to reach 15 characters.
    """
    if url.startswith("https://"):
        trimmed = url[len("https://") :]
    else:
        trimmed = url
    alphanum = re.sub(r"[^A-Za-z0-9]", "", trimmed)
    base = alphanum[:8]
    remaining = 15 - len(base)
    rand_chars = "".join(random.choices(string.ascii_letters + string.digits, k=remaining))
    return base + rand_chars

def convert_to_clear_english(text):
    """
    Convert the given text into clear, grammatical English using the OpenAI API.
    If the text exceeds the maximum token limit, it is truncated to that limit.
    """
    from config import OPENAI_CLIENT  # OPENAI_CLIENT is set to openai in config.py
    words = text.split()
    token_count = len(words)
    if token_count > MAX_TOKEN_LIMIT:
        st.warning(f"Token limit exceeded ({token_count} tokens > {MAX_TOKEN_LIMIT}). Truncating text for AI conversion.")
        text = " ".join(words[:MAX_TOKEN_LIMIT])
    prompt = f"Please rewrite the following text clearly:\n\n{text}\n\nRewrite:"
    try:
        # Instead of client.ChatCompletion.create(...), use openai.chat.completions.create(...)
        import openai
        response = openai.chat.completions.create(
            model="chatgpt-4o-latest",
            messages=[
                {"role": "system", "content": "Rewrite text in clear, concise English."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.5,
            max_tokens=16384
        )
        rewritten_text = response.choices[0].message.content.strip()
        return rewritten_text
    except Exception as e:
        st.error(f"OpenAI error: {str(e)}")
        return text

# New helper functions for recursive extraction consolidation

def recursive_extract_all(urls, max_depth=3):
    consolidated = {}
    for url in urls:
        data = web_scraper.recursive_extract(url, current_depth=0, max_depth=max_depth)
        consolidated[url] = data
    return consolidated

def format_recursive_data_for_word(data, indent=0):
    text = ""
    prefix = "    " * indent
    if isinstance(data, dict):
        text += f"{prefix}URL: {data.get('URL', '')}\n"
        text += f"{prefix}Title: {data.get('Title', '')}\n"
        text += f"{prefix}Headings: {data.get('Headings', {})}\n"
        text += f"{prefix}Content:\n{prefix}{data.get('FullContent', '')}\n"
        if data.get('sub_pages'):
            text += f"{prefix}Sub-pages:\n"
            for sub in data['sub_pages']:
                text += format_recursive_data_for_word(sub, indent+1)
    return text

def flatten_recursive_data(data, depth=0, rows=None):
    if rows is None:
        rows = []
    if isinstance(data, dict):
        rows.append({
            "URL": data.get("URL", ""),
            "Title": data.get("Title", ""),
            "Headings": str(data.get("Headings", {})),
            "Content": data.get("FullContent", ""),
            "Depth": depth
        })
        if data.get("sub_pages"):
            for sub in data["sub_pages"]:
                flatten_recursive_data(sub, depth+1, rows)
    return rows

def save_to_word(results, url):
    directory = "temp_uploaded_files"
    if not os.path.exists(directory):
        os.makedirs(directory)
    document = Document()
    document.add_heading('Consolidated Extraction Results', level=1)
    document.add_paragraph(f"Source URL: {url}")
    document.add_paragraph(results)
    filename = generate_random_filename(url) + ".docx"
    save_path = os.path.join(directory, filename)
    document.save(save_path)
    return save_path

def save_to_excel(rows, url):
    directory = "temp_uploaded_files"
    if not os.path.exists(directory):
        os.makedirs(directory)
    df = pd.DataFrame(rows)
    filename = generate_random_filename(url) + ".xlsx"
    save_path = os.path.join(directory, filename)
    df.to_excel(save_path, index=False)
    return save_path

def main():
    st.title("General Web Scraper")
    st.write("Enter URLs manually or upload a plain text file with URLs (one URL per line).")

    ########################################################################
    # NEW: Add a text area to capture user prompt or search instructions
    ########################################################################
    user_prompt = st.text_area(
        "Optionally describe specific info you want (e.g. names, titles, board members):",
        value="Please gather contact information for senior leadership, including the CEO and the Board of Directors, along with their names, titles, and contact details.",
        height=80,
    )
    ########################################################################

    uploaded_file = st.file_uploader("Upload a plain text file with URLs", type=["txt"])
    urls = []

    if uploaded_file is not None:
        try:
            file_content = io.StringIO(uploaded_file.read().decode("utf-8"))
        except Exception:
            st.error("Error reading the file. Please ensure it is a valid text file.")
            st.stop()
        file_urls = [line.strip() for line in file_content if line.strip()]
        if len(file_urls) > 100:
            st.error("Only 100 entries may be processed at a time. Please attach a file with 100 or fewer URLs or close.")
            st.stop()
        urls.extend(file_urls)
    else:
        manual_input = st.text_area("Enter one URL per line (press Enter after each URL):", "")
        if manual_input:
            manual_urls = [line.strip() for line in manual_input.splitlines() if line.strip()]
            if len(manual_urls) > 9:
                st.error("For more than 9 entries, please attach a plain text file with the full list.")
                st.stop()
            urls.extend(manual_urls)

    if not urls:
        st.error("No URLs provided. Please input or upload URLs and try again.")
        st.stop()

    # Normalize URLs
    normalized_urls = []
    for url in urls:
        norm_url = normalize_url(url)
        if norm_url:
            normalized_urls.append(norm_url)

    st.subheader("Normalized URLs:")
    cols = st.columns(4)
    for idx, url in enumerate(normalized_urls):
        col = cols[idx % len(cols)]
        col.write(url)

    ########################################################################
    # ENHANCED: If the user provided a prompt, perform enhanced site analysis
    ########################################################################
    if user_prompt.strip():
        st.write("Discovering site structure and analyzing content...")
        prioritized_list, ai_analysis = web_scraper.ai_assisted_robots_analysis(normalized_urls, user_prompt)
        urls_to_crawl = prioritized_list

        st.subheader("Identified Nodes for Crawl:")
        for node in urls_to_crawl:
            st.write(node)
            
        if ai_analysis and len(ai_analysis) > 0:
            st.subheader("AI Analysis:")
            for analysis in ai_analysis:
                processed_analysis = process_analysis_with_hyperlinks(analysis, normalized_urls[0] if normalized_urls else "")
                st.markdown(processed_analysis, unsafe_allow_html=True)
    else:
        urls_to_crawl = normalized_urls
    ########################################################################

    # NEW: Perform recursive extraction (depth = 3) on the AI-selected nodes
    with st.spinner("Performing recursive extraction (depth = 3)..."):
        consolidated_data = recursive_extract_all(urls_to_crawl, max_depth=3)
    
    # Prepare consolidated content for Word export
    word_content = ""
    for url, data in consolidated_data.items():
        word_content += format_recursive_data_for_word(data) + "\n\n"
    
    # Flatten the data for Excel export
    excel_rows = []
    for url, data in consolidated_data.items():
        excel_rows.extend(flatten_recursive_data(data))
    
    # Save consolidated output to Word and Excel files
    word_filepath = save_to_word(word_content, normalized_urls[0])
    excel_filepath = save_to_excel(excel_rows, normalized_urls[0])
    
    st.success("Extraction and export completed successfully!")
    with open(word_filepath, "rb") as f_word:
        st.download_button(
            label="📥 Download Word File",
            data=f_word,
            file_name=os.path.basename(word_filepath),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    with open(excel_filepath, "rb") as f_excel:
        st.download_button(
            label="📥 Download Excel File",
            data=f_excel,
            file_name=os.path.basename(excel_filepath),
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

def process_analysis_with_hyperlinks(analysis_text, base_url):
    """
    Process the analysis text to add hyperlinks to key sections.
    
    Args:
        analysis_text: The text from AI analysis
        base_url: The base URL of the website being analyzed
    
    Returns:
        Processed text with hyperlinks added
    """
    if base_url:
        parsed_url = urlparse(base_url)
        base_url = f"{parsed_url.scheme}://{parsed_url.netloc}"
    else:
        return analysis_text
    sections = {
        "About Us": "/about-us",
        "Contact Us": "/contact-us",
        "Vision, Mission, and Values": "/about-us/mission",
        "News & Resources": "/news"
    }
    processed_text = analysis_text
    for section, path in sections.items():
        full_url = urljoin(base_url, path)
        hyperlink = f'<a href="{full_url}" target="_blank">{section}</a>'
        pattern = r'(\d+\.\s*)?' + re.escape(section)
        processed_text = re.sub(pattern, r'\1' + hyperlink, processed_text)
    return processed_text

def run_app():
    main()

if __name__ == "__main__":
    run_app()
